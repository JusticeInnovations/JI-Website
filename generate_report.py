"""Generate Justice Innovations Website Technical & Functional Report as DOCX."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = r"C:\Users\kbatj\Desktop\JI\Website - JI\JI-Website-Report.docx"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(doc, text="", bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        cell._tc.get_or_add_tcPr().append(
            _shading_element("1B3A6B")
        )
    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=10)
            if ri % 2 == 1:
                cell._tc.get_or_add_tcPr().append(_shading_element("EEF2F7"))
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def _shading_element(fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    return shd

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B3A6B")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── build document ─────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("JUSTICE INNOVATIONS")
set_font(run, size=28, bold=True, color=(13, 31, 53))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Website Technical & Functional Report")
set_font(run, size=18, italic=True, color=(27, 111, 232))

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prosecutor-Guided Intake Technology")
set_font(run, size=13, color=(80, 80, 80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("justiceinnovations.github.io/JI-Website")
set_font(run, size=11, color=(27, 111, 232))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Report Date: {datetime.date.today().strftime('%B %d, %Y')}")
set_font(run, size=10, color=(120, 120, 120))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Confidential — Justice Innovations Internal Use")
set_font(run, size=10, color=(150, 150, 150))

doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────

add_heading(doc, "Table of Contents", 1)
toc_items = [
    ("1", "Executive Summary"),
    ("2", "Technology Stack & Architecture"),
    ("3", "Design System"),
    ("  3.1", "Color Tokens"),
    ("  3.2", "Typography"),
    ("  3.3", "Layout & Spacing"),
    ("  3.4", "Responsive Breakpoints"),
    ("  3.5", "Animation System"),
    ("4", "Information Architecture & Content"),
    ("  4.1", "Site Sections"),
    ("  4.2", "Navigation Structure"),
    ("  4.3", "Key Statistics & Data Points"),
    ("  4.4", "Publications & Research"),
    ("  4.5", "Compliance & Trust Signals"),
    ("5", "Interactive Features & JavaScript"),
    ("  5.1", "Modal System"),
    ("  5.2", "Contact Form"),
    ("  5.3", "Navigation Behavior"),
    ("  5.4", "Animations & Scroll Effects"),
    ("6", "Accessibility & WCAG Compliance"),
    ("7", "Security"),
    ("8", "Mobile Responsiveness"),
    ("9", "Testing Suite"),
    ("  9.1", "Test Coverage"),
    ("  9.2", "CI/CD Pipeline"),
    ("10", "SEO & Metadata"),
    ("11", "Known Limitations & Recommendations"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{num}{'.' if not num.startswith(' ') else '  '} {title}")
    set_font(run, size=10.5, bold=(not num.startswith(" ")))

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────

add_heading(doc, "1. Executive Summary", 1)
add_para(doc, (
    "The Justice Innovations marketing website is a single-page application (SPA) "
    "delivered as a single self-contained HTML file (index.html). It serves as the "
    "primary digital presence for Justice Innovations, a Houston, Texas-based Public "
    "Benefit Corporation that develops prosecutor-guided intake software for law "
    "enforcement agencies, district attorneys, courts, and legislators — primarily "
    "in Texas."
), space_after=8)
add_para(doc, (
    "The site presents the company's flagship product, Intake, which connects "
    "arresting officers to a prosecutor hotline before any booking is made. The "
    "platform enables a prosecutor to review the legal elements of a potential case "
    "in real time and return one of four decisions: Accept, Deny, Hold, or Needs "
    "More Information. The average call length is 90 seconds."
), space_after=8)
add_para(doc, (
    "The website is hosted on GitHub Pages from the main branch of the "
    "JusticeInnovations/JI-Website repository and goes live within 1–2 minutes of "
    "any commit. It requires no build tools, back-end servers, or frameworks. All "
    "styles (approximately 19 KB of CSS) and interactivity (approximately 350 lines "
    "of JavaScript) are embedded directly in index.html."
), space_after=8)
add_para(doc, "Key facts at a glance:", bold=True)
bullets = [
    "Single HTML file — no build system, no dependencies, no frameworks",
    "Hosted on GitHub Pages (public, no authentication)",
    "88 automated tests across 5 test files (pytest + BeautifulSoup4)",
    "CI/CD via GitHub Actions: black formatting, ruff linting, pytest on every push",
    "WCAG 2.1 AA compliant — skip links, ARIA roles, focus traps, focus-visible CSS",
    "3 font families: Cormorant Garamond (serif), Outfit (sans-serif), DM Mono (monospace)",
    "16 CSS custom property color tokens, 5 layout tokens",
    "23 named page sections with anchor IDs",
    "3 modal dialogs (legal docs, contact form, login placeholder)",
    "Form submissions via Formspree (https://formspree.io/f/xbdzzwqo)",
    "Social preview image (social-preview.svg) and full Open Graph/Twitter Card metadata",
    "JSON-LD structured data for Organization and SoftwareApplication schemas",
]
for b in bullets:
    add_bullet(doc, b)

doc.add_page_break()

# ── 2. TECHNOLOGY STACK ───────────────────────────────────────────────────────

add_heading(doc, "2. Technology Stack & Architecture", 1)
add_para(doc, (
    "The site is architected as a zero-dependency static page. All HTML structure, "
    "CSS styling, and JavaScript behavior are contained within a single index.html "
    "file of approximately 148 KB. The architectural decision to use a single file "
    "was intentional: it simplifies deployment (drag-and-drop or git push), "
    "eliminates asset pipeline complexity, and ensures the site remains portable and "
    "self-contained."
), space_after=8)

add_heading(doc, "Delivery & Hosting", 2)
add_table(doc,
    ["Component", "Technology", "Notes"],
    [
        ["Hosting", "GitHub Pages", "Served from main branch; deploys in 1–2 min"],
        ["Repository", "github.com/JusticeInnovations/JI-Website", "Public repository"],
        ["HTML", "HTML5 (single file)", "DOCTYPE html, lang=en, semantic elements"],
        ["CSS", "Vanilla CSS (embedded)", "~19 KB inline in <style> block"],
        ["JavaScript", "Vanilla ES6+ (embedded)", "~350 lines inline before </body>"],
        ["Fonts", "Google Fonts CDN", "Cormorant Garamond, Outfit, DM Mono"],
        ["Form Backend", "Formspree", "HTTPS POST to formspree.io/f/xbdzzwqo"],
        ["Version Control", "Git / GitHub", "main branch, no feature branches"],
        ["CI/CD", "GitHub Actions", ".github/workflows/ci.yaml"],
        ["Testing", "pytest + BeautifulSoup4", "Python 3.11, 88 tests"],
    ],
    col_widths=[1.5, 2.2, 3.0]
)

add_heading(doc, "File Structure", 2)
add_table(doc,
    ["File / Directory", "Purpose"],
    [
        ["index.html", "Entire website — HTML, CSS, JavaScript in one file (~148 KB)"],
        ["social-preview.svg", "1200×630 Open Graph image for social sharing cards"],
        ["requirements-test.txt", "Python test dependencies (pytest, beautifulsoup4, etc.)"],
        ["pyproject.toml", "pytest, black, and ruff configuration"],
        [".github/workflows/ci.yaml", "GitHub Actions CI pipeline definition"],
        ["tests/conftest.py", "pytest fixtures: html_path, soup, raw_html"],
        ["tests/test_content.py", "32 tests — copy, statistics, sections, navigation"],
        ["tests/test_accessibility.py", "24 tests — WCAG, ARIA, focus, modals, CSS"],
        ["tests/test_mobile.py", "8 tests — viewport, breakpoints, touch targets"],
        ["tests/test_security.py", "6 tests — noopener, form HTTPS, no credentials"],
        ["tests/test_html_structure.py", "18 tests — DOCTYPE, meta, headings, structure"],
        ["README.md", "Project overview, deployment instructions, badge links"],
        [".gitignore", "Excludes __pycache__, .pytest_cache, .ruff_cache"],
        ["JI-Website-Report.docx", "This document"],
    ],
    col_widths=[2.5, 4.2]
)

doc.add_page_break()

# ── 3. DESIGN SYSTEM ──────────────────────────────────────────────────────────

add_heading(doc, "3. Design System", 1)
add_para(doc, (
    "The design system is implemented entirely through CSS custom properties "
    "(CSS variables) defined in the :root selector. All colors, spacing, and layout "
    "measurements reference these tokens, ensuring visual consistency and making "
    "global changes trivial."
))

add_heading(doc, "3.1 Color Tokens", 2)
add_table(doc,
    ["Variable", "Value", "Usage"],
    [
        ["--white",      "#FFFFFF",  "Pure white text"],
        ["--blue",       "#0D1F35",  "Primary navy — section backgrounds, nav"],
        ["--blue-dark",  "#081628",  "Darker navy — footer background"],
        ["--off",        "#112540",  "Off-navy — card backgrounds, section--light"],
        ["--off-dark",   "#080F1C",  "Darkest background — section--dark"],
        ["--off-mid",    "#0F1923",  "Mid-dark — modal backgrounds"],
        ["--ink",        "#E8EDF5",  "Light text on dark backgrounds"],
        ["--electric",   "#1B6FE8",  "Primary brand blue — CTA buttons, highlights"],
        ["--electric-lt","#3D8EFF",  "Lighter blue — h1 em text, hover states"],
        ["--teal-lt",    "#00C2D4",  "Teal accent — hover borders, teal labels"],
        ["--gold",       "#D4860A",  "Gold — research card labels, timeline phases"],
        ["--gold-lt",    "#F5A623",  "Light gold — gold sec-label variant"],
        ["--red-lt",     "#FF6B6B",  "Coral red — section labels, eyebrow text, logo sub"],
        ["--green",      "#28C840",  "Status green — mockup indicators, tag.t-ok"],
        ["--amber",      "#FFBD2E",  "Amber — warning status, tag.t-warn"],
        ["--red-status", "#FF5F57",  "Status red — tag.t-flag, mockup dot"],
    ],
    col_widths=[1.5, 1.2, 4.0]
)

add_heading(doc, "3.2 Typography", 2)
add_para(doc, "Three font families are loaded from Google Fonts CDN:", space_after=4)
add_table(doc,
    ["Font Family", "Weights Loaded", "Usage"],
    [
        ["Cormorant Garamond", "500, 600, italic 500", "H1, H2, stat numbers, paper titles, logo name, footer brand"],
        ["Outfit", "300, 400, 500, 600", "Body text, navigation links, buttons, form inputs, intro paragraphs"],
        ["DM Mono", "400", "Section labels, eyebrow text, monospace tags, mockup UI, footer column labels"],
    ],
    col_widths=[1.8, 1.6, 3.3]
)

add_para(doc, "")
add_para(doc, "Key type scale:", bold=True, space_after=4)
add_table(doc,
    ["Element", "Size", "Font", "Weight", "Notes"],
    [
        ["H1", "clamp(52px, 6vw, 82px)", "Cormorant Garamond", "500", "Hero headline; em children use --electric-lt"],
        ["H2", "clamp(34px, 3.8vw, 52px)", "Cormorant Garamond", "500", "Section headings"],
        ["Nav logo name", "72px (desktop)", "Cormorant Garamond", "600", "Scales to 38px mobile"],
        ["Nav logo sub", "18px", "DM Mono", "400", "Uppercase, red-lt"],
        [".sec-label", "14px", "DM Mono", "400", "Uppercase, letter-spacing .16em"],
        [".hero-eyebrow", "14px", "DM Mono", "400", "Uppercase, letter-spacing .16em"],
        [".intro-p", "18px", "Outfit", "400", "opacity .82, line-height 1.75"],
        [".stat-num", "clamp(36px, 4vw, 52px)", "Cormorant Garamond", "500", "Statistics display numbers"],
        [".p-title", "21px", "Cormorant Garamond", "500", "Publication paper titles"],
        [".p-authors", "14px", "Outfit", "400", "opacity .62"],
        [".p-journal", "13px", "Outfit", "400", "italic, opacity .62"],
        [".rc-title", "24px", "Cormorant Garamond", "500", "How It Works step titles"],
        [".rc-body", "14px", "Outfit", "400", "opacity .55, line-height 1.7"],
        ["Body / buttons", "14px", "Outfit", "400–600", "General UI text"],
    ],
    col_widths=[1.5, 1.8, 1.8, 0.8, 2.0]
)

add_heading(doc, "3.3 Layout & Spacing", 2)
add_table(doc,
    ["Token", "Default", "Purpose"],
    [
        ["--px", "48px", "Horizontal padding — all sections and nav"],
        ["--py", "96px", "Vertical padding — all sections"],
        ["--mw", "1280px", "Max content width — .inner container"],
        ["--nh", "116px", "Nav bar height — used for scroll-padding-top and position offsets"],
        ["--r",  "10px",  "Standard border-radius for cards, modals"],
    ],
    col_widths=[1.0, 1.0, 4.7]
)

add_heading(doc, "3.4 Responsive Breakpoints", 2)
add_table(doc,
    ["Breakpoint", "Trigger", "--px", "--py", "--nh", "Key Changes"],
    [
        ["Mobile S", "max-width: 479px", "16px", "48px", "—", "Smallest text overrides; h1 clamp(28px,6vw,44px)"],
        ["Mobile", "max-width: 767px", "20px", "64px", "80px", "Hamburger visible; nav links/actions hidden; all grids collapse to 1 col; h1 clamp(36px,10vw,52px)"],
        ["Tablet", "768px–1023px", "32px", "80px", "100px", "2-col grids active; logo name 52px; some 3/4-col grids become 2-col"],
        ["Desktop", "1024px+", "48px", "96px", "116px", "All multi-col grids; logo name 72px; full nav"],
        ["Large", "1600px+", "80px", "120px", "—", "h1 clamp(64px,5vw,96px); h2 clamp(40px,3.5vw,64px)"],
        ["XL", "1920px+", "120px", "140px", "140px", "Maximum padding; largest nav height"],
    ],
    col_widths=[0.9, 1.4, 0.5, 0.5, 0.5, 2.8]
)

add_heading(doc, "3.5 Animation System", 2)
add_para(doc, "The site uses four CSS keyframe animations plus CSS transitions on interactive elements:")
add_table(doc,
    ["Animation", "Duration", "Used On", "Description"],
    [
        ["fadeUp", "0.6s ease", "Hero eyebrow, H1, hero-p, hero-btns (staggered 0/80/160/220ms)", "Opacity 0→1, translateY 16px→0 on page load"],
        ["flashBlue", "2.5s infinite", ".pl-blue, .pl-blue2 (offset 1.25s)", "Police light blue flash — opacity pulses 0.9→0.05→0.85→0.05"],
        ["flashRed", "2.5s infinite", ".pl-red, .pl-red2 (offset 1.25s)", "Police light red flash — alternates with flashBlue"],
        ["reflectShimmer", "2.5s infinite", ".pl-reflect", "Ground reflection shimmer synced with lights"],
        ["streakFlash", "2.5s infinite", ".pl-streak", "Horizontal light streak across bottom of hero"],
    ],
    col_widths=[1.4, 1.0, 2.1, 2.2]
)
add_para(doc, (
    "All animations are disabled when the user has prefers-reduced-motion: reduce "
    "set in their OS accessibility settings. The .police-lights element is hidden "
    "entirely (display: none) in the reduced-motion media query block."
), space_after=6)

doc.add_page_break()

# ── 4. INFORMATION ARCHITECTURE ───────────────────────────────────────────────

add_heading(doc, "4. Information Architecture & Content", 1)

add_heading(doc, "4.1 Site Sections", 2)
add_para(doc, (
    "The page is organized into 23 named sections, each with a unique anchor ID "
    "for deep linking and navigation. Sections alternate between dark "
    "(section--dark: #080F1C) and light (section--light: #112540) backgrounds to "
    "create visual rhythm."
))
add_table(doc,
    ["#ID", "Theme", "Heading / Purpose"],
    [
        ["top", "Dark (hero)", "\"The right call. At the right time.\" — Animated hero with police lights"],
        ["what-is-intake", "Dark", "\"The arrest decision is the most consequential moment...\" — Product overview; officer quote; $4.6M stat"],
        ["why-prosecutor", "Light", "\"Officers arrest. Prosecutors decide if charges hold.\" — Methodology rationale"],
        ["how-it-works", "Dark", "4-step process walkthrough (Officer → Prosecutor → Sheriff → Court Clerk)"],
        ["features", "Light", "\"Everything the process needs. Nothing it doesn't.\" — 10 feature cards"],
        ["implementation", "Dark", "\"A complete operational transition.\" — 3-phase timeline (Onboard, Train, Go Live)"],
        ["what-we-measure", "Light", "\"Outcomes that matter to the system — and the people in it.\" — 6 impact metrics"],
        ["research", "Dark", "\"The science behind the software.\" — Harris County Model; 45 years of proof"],
        ["research-data", "Light", "\"Stop managing cases that should never have been filed.\" — Cost analysis stats"],
        ["research-papers", "Dark", "\"The science.\" — 6 publications + white paper"],
        ["solutions", "Gradient", "\"The same problem. Four different perspectives.\" — Role overview"],
        ["solutions-da", "Dark", "District Attorneys & Prosecutors — role-specific benefits"],
        ["solutions-le", "Light", "Law Enforcement — role-specific benefits"],
        ["solutions-courts", "Dark", "Courts & Judges — role-specific benefits"],
        ["solutions-policy", "Light", "Legislators & Policy — role-specific benefits"],
        ["about", "Gradient", "\"Built for the system — and everyone it touches.\" — Company mission"],
        ["about-scilaw", "Dark", "SciLaw (Center for Science and Law) research partnership"],
        ["about-pilot", "Light", "Cameron County, TX pilot program details"],
        ["cta", "Light", "\"Ready when you are.\" — Final demo request CTA"],
    ],
    col_widths=[1.5, 1.0, 4.2]
)

add_heading(doc, "4.2 Navigation Structure", 2)
add_para(doc, (
    "The fixed navigation bar contains a logo, five dropdown menus, a Login button, "
    "and a Talk to Us CTA. On mobile (≤767px), nav links and actions are hidden and "
    "replaced by a hamburger button that opens a full-width drawer."
))
add_para(doc, "Desktop dropdown menus:", bold=True, space_after=3)
nav_items = [
    ("Home", "#top", "—"),
    ("Intake", "#what-is-intake", "What Is Intake?, How It Works, Why a Prosecutor?, Features & Platform, Implementation Package, What We Measure, Request a Demo"),
    ("Research", "#research", "The Harris County Model, The Cost of Doing Nothing, Published Papers, White Papers, Fairness & Equity Framework"),
    ("Our Technology", "#solutions", "District Attorneys & Prosecutors, Law Enforcement, Courts & Judges, Legislators & Policy"),
    ("About", "#about", "Who We Are, SciLaw Partnership, Cameron County Pilot, Contact Us"),
]
add_table(doc,
    ["Menu Item", "Primary href", "Dropdown Items"],
    nav_items,
    col_widths=[1.3, 1.5, 3.9]
)

add_heading(doc, "4.3 Key Statistics & Data Points", 2)
add_table(doc,
    ["Statistic", "Value", "Context / Section"],
    [
        ["Annual savings", "$4.6M", "Cameron County, TX — 25% reduction in DUI/DV pretrial population"],
        ["Hotline call length", "90 seconds", "Average officer-to-prosecutor decision time"],
        ["No Action jail nights", "36 avg.", "Average nights a No Action case spends in jail without screening"],
        ["Pretrial reduction", "25%", "Average reduction when pre-arrest screening adopted"],
        ["Releases (Cameron)", "212 people", "Projected releases with 25% reduction in Cameron County"],
        ["Harris County model", "Since 1977", "Uninterrupted operation of pre-arrest screening"],
        ["Miami-Dade waste", "$20M+", "Projected annual waste Intake could have prevented (2000–2011)"],
        ["Texas statewide savings", "$238M", "Modeled annual savings with 25% pretrial reduction (10,883 people)"],
        ["Bexar County savings", "$15.4M", "Annual savings for Bexar County alone — 707 people released"],
        ["U.S. county ranking", "Top 25", "Harris County: highest arrest-to-prosecution rate among 25 largest U.S. counties"],
    ],
    col_widths=[1.8, 1.0, 3.9]
)

add_heading(doc, "4.4 Publications & Research", 2)
add_para(doc, (
    "The Publications section is organized into three groups. All publications "
    "are linked to external sources where available and open in a new tab "
    "with rel=\"noopener\" for security."
))
add_para(doc, "Group 1 — Core Research (gold accent):", bold=True, space_after=3)
add_table(doc,
    ["Title", "Authors", "Journal", "Year"],
    [
        ["Quantifying the Prosecutorial Preauthorization Intake System and the Costs of 'No Action' Cases",
         "Ormachea, Davenport, Haarsma & Eagleman · CSL", "Journal of Science and Law", "2019"],
        ["A New Criminal Records Database for Large-Scale Analysis of Policy and Behavior",
         "Ormachea, Haarsma, Davenport & Eagleman · CSL", "Journal of Science and Law", "2015"],
    ],
    col_widths=[2.5, 1.8, 1.5, 0.6]
)
add_para(doc, "Group 2 — Scientific Foundation (teal accent):", bold=True, space_after=3)
add_table(doc,
    ["Title", "Authors", "Journal", "Year"],
    [
        ["Assessing Risk Among Correctional Community Probation Populations",
         "Haarsma, Davenport et al. · CSL", "Frontiers in Psychology", "2019"],
        ["Enabling Individualized Criminal Sentencing While Reducing Subjectivity",
         "Ormachea, Davenport, Haarsma et al. · CSL", "AMA Journal of Ethics", "2016"],
        ["The Role of Tablet-Based Psychological Tasks in Risk Assessment",
         "Ormachea, Lovins, Eagleman et al. · CSL", "Criminal Justice and Behavior · SAGE", "2017"],
        ["Gendered Outcomes in Prostitution Arrests in Houston, Texas",
         "Pfeffer, Ormachea & Eagleman · CSL", "Crime & Delinquency · SAGE", "2017"],
    ],
    col_widths=[2.5, 1.8, 1.5, 0.6]
)
add_para(doc, "Group 3 — Technical Documentation (white accent):", bold=True, space_after=3)
add_table(doc,
    ["Document", "Authors", "Availability"],
    [
        ["Justice Innovations Intake Platform: Security & Compliance White Paper",
         "Justice Innovations Technical Team", "Download (JI-Security-Compliance-White-Paper.pdf)"],
        ["Scientific Foundations of the Intake Decision Framework: Evidence, Bias, and Accountability",
         "Justice Innovations Research Team · CSL", "Available on Request"],
    ],
    col_widths=[2.8, 2.0, 1.9]
)

add_heading(doc, "4.5 Compliance & Trust Signals", 2)
add_para(doc, "The trust strip below the hero displays seven compliance badges with hover tooltips:")
add_table(doc,
    ["Badge", "Full Name", "Tooltip Description"],
    [
        ["CJIS-Aligned Architecture", "Criminal Justice Information Services", "FBI security policy governing all criminal justice data systems"],
        ["FedRAMP-Ready Infrastructure", "Federal Risk & Authorization Management Program", "Hosted on AWS GovCloud, which holds FedRAMP authorization"],
        ["TX-RAMP Pathway", "Texas Risk and Authorization Management Program", "Texas state cloud security certification pathway"],
        ["TIBRS / NIBRS Ready", "Texas/National Incident-Based Reporting System", "Integrated for automated case reporting"],
        ["Texas CCP Aligned", "Texas Code of Criminal Procedure", "Case workflow aligned to Texas CCP requirements"],
        ["Pilot-Ready", "—", "Structured for rigorously designed pilot evaluation"],
        ["WCAG 2.1 AA", "Web Content Accessibility Guidelines", "Accessibility standard compliance"],
    ],
    col_widths=[1.8, 1.9, 3.0]
)
add_para(doc, "")
add_para(doc, "Footer credential badges (bottom of page):", bold=True, space_after=3)
for b in [
    "WOSB — Woman-Owned Small Business (federally certified)",
    "Public Benefit Corporation — incorporated under Texas PBC statute",
    "Made in Texas",
    "CJIS Compliant",
    "AWS GovCloud Hosted",
]:
    add_bullet(doc, b)

doc.add_page_break()

# ── 5. INTERACTIVE FEATURES ───────────────────────────────────────────────────

add_heading(doc, "5. Interactive Features & JavaScript", 1)
add_para(doc, (
    "All JavaScript is vanilla ES6+ embedded in a single <script> block "
    "immediately before the closing </body> tag. There are no external JavaScript "
    "dependencies, no frameworks, and no module bundlers."
))

add_heading(doc, "5.1 Modal System", 2)
add_para(doc, (
    "Three modal dialogs are implemented. All modals follow WCAG 2.1 AA requirements: "
    "they carry role=\"dialog\", aria-modal=\"true\", and aria-labelledby referencing "
    "their visible title. A focus trap prevents keyboard users from tabbing outside "
    "the modal while it is open. Pressing Escape closes any open modal."
))
add_table(doc,
    ["Modal ID", "Purpose", "Trigger(s)", "ARIA labelledby"],
    [
        ["contactModal", "Contact / demo request form", "Nav CTA, hero button, solution CTAs, footer CTA, mobile nav CTA", "contactModalTitle"],
        ["legalModal", "Legal documents viewer (Privacy, Terms, Accessibility, DPA)", "Footer links for each document", "legalContent"],
        ["loginModal", "Login portal placeholder (coming soon)", "Nav Login button", "loginModalTitle"],
    ],
    col_widths=[1.4, 2.4, 2.2, 1.7]
)

add_heading(doc, "5.2 Contact Form", 2)
add_para(doc, "The contact form within contactModal collects the following fields:")
add_table(doc,
    ["Field ID", "Type", "Label", "Required", "Placeholder"],
    [
        ["m-name",    "text",     "Name",                 "Yes", "Your name"],
        ["m-email",   "email",    "Email",                "Yes", "you@example.com"],
        ["m-org",     "text",     "Organization / County","No",  "e.g. Cameron County DA"],
        ["m-phone",   "tel",      "Phone",                "No",  "(000) 000-0000"],
        ["m-message", "textarea", "Message",              "Yes", "Tell us what you're working on..."],
    ],
    col_widths=[1.0, 0.8, 1.8, 0.8, 2.3]
)
add_para(doc, "")
add_para(doc, (
    "On submit, the JavaScript validates all required fields client-side, then "
    "POSTs the data as JSON to https://formspree.io/f/xbdzzwqo. On success, the "
    "form is hidden and a confirmation message is displayed inside the modal: "
    "\"Message Sent — We'll be in touch within 24 hours.\" The modal title "
    "updates dynamically based on which trigger was used (\"Talk to Us\", "
    "\"Request Demo\", etc.)."
))

add_heading(doc, "5.3 Navigation Behavior", 2)
add_para(doc, "Desktop navigation dropdowns:", bold=True, space_after=3)
add_para(doc, (
    "Each top-level nav item with a subnav has mouseenter/mouseleave event "
    "listeners with a 300 ms debounce timeout. Hovering adds the .open class to "
    "the .subnav child, which transitions it from opacity 0 / translateY(-6px) "
    "to opacity 1 / translateY(0). Hovering over the subnav itself prevents it "
    "from closing prematurely."
), space_after=6)
add_para(doc, "Mobile hamburger menu:", bold=True, space_after=3)
add_para(doc, (
    "Clicking the hamburger button (#hamburgerBtn) toggles the .is-open class on "
    "#mobileNav (display: none → display: block). The button's aria-expanded "
    "attribute is toggled between \"false\" and \"true\", and aria-label switches "
    "between \"Open menu\" and \"Close menu\". Each mobile nav link calls "
    "closeMnav() on click to dismiss the drawer after navigation."
))

add_heading(doc, "5.4 Animations & Scroll Effects", 2)
add_para(doc, (
    "An IntersectionObserver watches all elements with the .au class "
    "(used as a general animation utility throughout the page). When an .au "
    "element enters the viewport, a CSS class is added to trigger its entrance "
    "animation. This creates staggered reveal effects as the user scrolls "
    "through sections."
), space_after=6)
add_para(doc, (
    "The hero section police lights animation runs continuously using four "
    "CSS keyframe animations (flashBlue, flashRed, reflectShimmer, streakFlash) "
    "on absolutely-positioned radial gradient divs. Blue and red lights are "
    "offset by 1.25 seconds to create a realistic alternating police strobe effect."
))

doc.add_page_break()

# ── 6. ACCESSIBILITY ──────────────────────────────────────────────────────────

add_heading(doc, "6. Accessibility & WCAG Compliance", 1)
add_para(doc, (
    "The site targets WCAG 2.1 Level AA compliance. Accessibility was a primary "
    "design consideration and is verified by 24 automated tests in "
    "tests/test_accessibility.py."
))
add_table(doc,
    ["Requirement", "Implementation", "Test"],
    [
        ["Skip navigation", "<a class=\"skip-link\" href=\"#main-content\"> — reveals on :focus", "test_skip_link_exists, test_skip_link_targets_main, test_skip_link_is_first_focusable"],
        ["Main landmark", "<main id=\"main-content\"> wraps all page content", "Verified by skip link target"],
        ["Modal role", "All modals: role=\"dialog\" aria-modal=\"true\"", "test_all_modals_have_dialog_role, test_all_modals_have_aria_modal"],
        ["Modal labels", "All modals: aria-labelledby referencing visible title ID", "test_all_modals_have_aria_labelledby"],
        ["Modal close buttons", "class=\"modal-close\" with aria-label=\"Close\"", "test_modal_close_buttons_have_aria_label"],
        ["Focus trap", "trapFocus() / releaseFocus() JavaScript functions", "Behavioral — not static-tested"],
        ["Hamburger ARIA", "aria-label, aria-expanded, aria-controls attributes", "test_hamburger_has_aria_label/expanded/controls"],
        ["Semantic buttons", "All modal triggers and nav CTAs are <button>, not <a href='#'>", "test_no_anchor_onclick_hash, test_nav_cta_is_button, test_nav_login_is_button"],
        ["Form labels", "All inputs and textareas have associated <label for=...>", "test_form_inputs_have_labels, test_textarea_has_label"],
        ["Decorative SVGs", "All decorative SVGs carry aria-hidden=\"true\"", "test_logo_svg_aria_hidden, test_feature_card_svgs_aria_hidden, test_impact_card_svgs_aria_hidden"],
        ["Focus visibility", ":focus-visible CSS defined for all interactive elements", "test_focus_visible_styles_defined"],
        ["Reduced motion", "prefers-reduced-motion: reduce hides police lights, disables animations", "test_prefers_reduced_motion_defined, test_police_lights_hidden_on_reduced_motion"],
        ["No zoom blocking", "Viewport meta does not include user-scalable=no or maximum-scale=1", "test_viewport_meta_no_user_scalable_no"],
        ["Color contrast", "--red-lt #FF6B6B on navy #0D1F35 = 5.98:1 (passes AA)", "Manual audit"],
        ["Footer links", "Opacity raised to .72 for WCAG contrast compliance", "Manual audit"],
    ],
    col_widths=[1.7, 2.5, 2.5]
)

doc.add_page_break()

# ── 7. SECURITY ───────────────────────────────────────────────────────────────

add_heading(doc, "7. Security", 1)
add_para(doc, (
    "As a static site, the attack surface is limited. Security best practices "
    "are verified by 6 automated tests in tests/test_security.py."
))
add_table(doc,
    ["Control", "Implementation", "Test"],
    [
        ["External link safety", "All target=\"_blank\" links have rel=\"noopener\"", "test_external_links_have_noopener"],
        ["No javascript: hrefs", "No anchor tags use javascript: protocol", "test_no_javascript_hrefs"],
        ["No onclick anchors", "Modal triggers are <button> elements, not <a href='#' onclick>", "test_no_inline_event_handlers_on_anchors"],
        ["HTTPS form endpoint", "Formspree endpoint uses https://formspree.io/f/xbdzzwqo", "test_form_uses_https_endpoint"],
        ["No credentials in HTML", "Scanned for password=, api_key=, secret=, aws_access_key patterns", "test_no_sensitive_data_in_html"],
        ["Formspree ID format", "Formspree form ID xbdzzwqo is ≥6 alphanumeric characters", "test_formspree_id_is_expected_format"],
    ],
    col_widths=[1.8, 2.7, 2.2]
)
add_para(doc, "")
add_para(doc, "Data Processing & CJIS Compliance (documented in DPA modal):", bold=True, space_after=3)
for item in [
    "All CJI in transit encrypted with TLS 1.2+",
    "All CJI at rest encrypted with AES-256",
    "Hosted exclusively on AWS GovCloud (US-East and US-West)",
    "Role-based access control (RBAC) with 4 defined roles: Officer, ADA/Prosecutor, Sheriff/Jail Intake, Court Clerk",
    "MFA required for all users; sessions expire on inactivity",
    "Immutable audit logs with timestamps, user IDs, and action descriptions",
    "Fingerprint-based background checks for JI personnel with CJI access",
    "Incident response plan — agency notification within 72 hours of confirmed breach",
    "Signed Data Processing Addendum required before any CJI transmission",
]:
    add_bullet(doc, item)

doc.add_page_break()

# ── 8. MOBILE RESPONSIVENESS ──────────────────────────────────────────────────

add_heading(doc, "8. Mobile Responsiveness", 1)
add_para(doc, (
    "Mobile responsiveness is verified by 8 automated tests in tests/test_mobile.py. "
    "The site uses a mobile-first approach with CSS custom properties that scale "
    "at defined breakpoints."
))
add_table(doc,
    ["Feature", "Desktop Behavior", "Mobile Behavior (≤767px)"],
    [
        ["Navigation", "Full horizontal nav with dropdowns", "Hidden; replaced by hamburger + drawer (#mobileNav)"],
        ["Hero H1", "clamp(52px, 6vw, 82px)", "clamp(36px, 10vw, 52px)"],
        ["Hero buttons", "Inline row", "Full-width column stack"],
        ["Logo name", "72px Cormorant Garamond", "38px"],
        ["Logo icon", "96×96px", "54×54px"],
        ["Feature grid", "auto-fit minmax(240px, 1fr)", "Single column"],
        [".g2 / .g3 / .g4 grids", "2/3/4 columns", "Single column (1fr override)"],
        ["Impact grid", "auto-fit minmax(240px, 1fr)", "Single column"],
        ["Stat banner", "4 columns", "1 column stacked"],
        ["Status flow mockup", "Full width", "overflow-x: auto scroll container"],
        ["Trust strip", "Horizontal flex row", "Vertical column stack"],
        ["Research papers", "2-column grid", "Single column"],
        ["H1 font size", "Responsive clamp", "Verified ≥ 28px minimum"],
        ["Viewport meta", "width=device-width, initial-scale=1.0", "No user-scalable=no; pinch-zoom allowed"],
        ["CTA button (mobile nav)", "N/A", ".mnav__cta button, full-width, 15px"],
    ],
    col_widths=[1.7, 2.3, 2.7]
)

doc.add_page_break()

# ── 9. TESTING SUITE ──────────────────────────────────────────────────────────

add_heading(doc, "9. Testing Suite", 1)
add_para(doc, (
    "The test suite consists of 88 pytest tests across 5 test files. Tests use "
    "BeautifulSoup4 to parse the rendered HTML of index.html and assert structural, "
    "content, accessibility, mobile, and security properties. Tests are configured "
    "in pyproject.toml and run automatically on every push to main via GitHub Actions."
))

add_heading(doc, "9.1 Test Coverage", 2)
add_table(doc,
    ["Test File", "Tests", "Coverage Area"],
    [
        ["test_html_structure.py", "18", "DOCTYPE, lang attribute, charset, viewport meta, title, meta description, OG tags, favicon, H1, H2 presence, nav element, main element, footer, internal anchors, external noopener links, Google Fonts"],
        ["test_accessibility.py", "24", "Skip link (exists, target, DOM order), modals (dialog role, aria-modal, aria-labelledby, close label), hamburger (aria-label, aria-expanded, aria-controls), no onclick anchors, nav CTA/login as buttons, form labels (inputs + textarea), SVG aria-hidden (logo, fc-icon, impact-icon), :focus-visible CSS, prefers-reduced-motion, skip link :focus CSS, police-lights display:none in reduced motion"],
        ["test_content.py", "32", "All David Eagleman copy edits (contrastive constructions removed, specific phrasings), key statistics ($4.6M, 25%, 90s, 36, 1977), pilot agencies (Cameron County, Brownsville), offense types (DUI, Domestic Violence), WOSB/PBC certifications, required section IDs, contact form, white paper link, navigation structure (Intake, Research, Talk to Us button)"],
        ["test_mobile.py", "8", "No user-scalable=no in viewport, @media breakpoints defined, max-width 400/700px queries, #mobileNav exists, #hamburgerBtn is button, H1 uses clamp(), .g2/.g3/.g4 have 1fr collapse, .mnav__cta is button, overflow-x: auto present"],
        ["test_security.py", "6", "target=_blank links have noopener, no javascript: hrefs, no onclick anchors, Formspree HTTPS endpoint, no sensitive data patterns, Formspree ID format validation"],
    ],
    col_widths=[1.8, 0.6, 4.3]
)

add_heading(doc, "9.2 CI/CD Pipeline", 2)
add_para(doc, (
    "The CI pipeline runs on every push to main and every pull request targeting main. "
    "It uses a single job (Lint & Test) on ubuntu-latest with Python 3.11."
))
add_table(doc,
    ["Step", "Tool", "Command", "Purpose"],
    [
        ["1 — Checkout", "actions/checkout@v4", "—", "Clone repository"],
        ["2 — Setup Python", "actions/setup-python@v5", "python-version: 3.11, cache: pip", "Install Python, restore pip cache"],
        ["3 — Install deps", "pip", "pip install -r requirements-test.txt", "Install beautifulsoup4, html5lib, lxml, pytest, black, ruff"],
        ["4 — Format check", "Black", "black --check tests/", "Fail if any test file is not Black-formatted"],
        ["5 — Lint", "Ruff", "ruff check tests/", "Enforce E/W/F/I/B/UP/ANN/S/RUF rules (S101 ignored for pytest asserts; ANN001/ANN201 ignored in tests)"],
        ["6 — Test", "pytest", "pytest tests/ -v --tb=short", "Run all 88 tests; verbose output; short tracebacks"],
        ["7 — Upload artifacts", "actions/upload-artifact@v4", ".pytest_cache/ retained 7 days", "Preserve test results for debugging"],
    ],
    col_widths=[1.5, 1.5, 2.2, 1.5]
)

add_para(doc, "")
add_para(doc, "Test dependencies (requirements-test.txt):", bold=True, space_after=3)
for dep in ["beautifulsoup4", "html5lib", "lxml", "pytest", "black", "ruff"]:
    add_bullet(doc, dep)

doc.add_page_break()

# ── 10. SEO & METADATA ────────────────────────────────────────────────────────

add_heading(doc, "10. SEO & Metadata", 1)
add_table(doc,
    ["Tag", "Value"],
    [
        ["<title>", "Justice Innovations — Prosecutor-Guided Intake Technology"],
        ["meta description", "Justice Innovations builds prosecutor-guided intake software that brings prosecutorial review to the moment of arrest..."],
        ["og:title", "Justice Innovations — Prosecutor-Guided Intake Technology"],
        ["og:description", "Evidence-based justice technology. Prosecutor-guided software that brings scientific rigor..."],
        ["og:type", "website"],
        ["og:url", "https://justiceinnovations.github.io/JI-Website/"],
        ["og:image", "https://justiceinnovations.github.io/JI-Website/social-preview.png"],
        ["og:image:width / height", "1200 / 630"],
        ["twitter:card", "summary_large_image"],
        ["twitter:title", "Justice Innovations — Prosecutor-Guided Intake Technology"],
        ["twitter:image", "https://justiceinnovations.github.io/JI-Website/social-preview.png"],
        ["canonical", "https://justiceinnovations.github.io/JI-Website/"],
        ["viewport", "width=device-width, initial-scale=1.0"],
        ["charset", "UTF-8"],
        ["lang", "en"],
        ["favicon", "Inline SVG data URI — navy rect with gold scale of justice icon"],
    ],
    col_widths=[2.0, 4.7]
)

add_para(doc, "")
add_heading(doc, "JSON-LD Structured Data", 2)
add_para(doc, "Two schema.org entities are defined in a <script type=\"application/ld+json\"> block in the <head>:")
add_table(doc,
    ["@type", "Key Fields"],
    [
        ["Organization", "name, url, description, foundingLocation (Houston TX), areaServed (US), keywords"],
        ["SoftwareApplication", "name (Intake), applicationCategory (BusinessApplication), operatingSystem (Web), description, offers, provider (@id reference to Organization)"],
    ],
    col_widths=[2.0, 4.7]
)

doc.add_page_break()

# ── 11. KNOWN LIMITATIONS ─────────────────────────────────────────────────────

add_heading(doc, "11. Known Limitations & Recommendations", 1)

add_heading(doc, "Content Gaps", 2)
recs_content = [
    ("No pricing information", "High", "Government buyers need cost signals. Add a pricing section or explicit \"contact for pricing\" prompt."),
    ("No team/leadership bios", "High", "Founder credibility is critical for government procurement decisions. Add a leadership section."),
    ("No customer testimonials", "High", "Cameron County is mentioned but no direct quotes from prosecutors or officers."),
    ("No FAQ section", "Medium", "Common procurement questions (implementation timeline, data ownership, integration, downtime) are unanswered."),
    ("No press/media mentions", "Medium", "An \"As seen in\" or press section builds credibility with new visitors."),
    ("Role sections are text-heavy", "Medium", "Four nearly identical role sections (DA, LE, Courts, Legislators) could be consolidated into a tabbed interface."),
    ("$4.6M stat not prominent enough", "Low", "The most compelling financial number is currently in a data-bar component below the officer quote."),
]
add_table(doc,
    ["Issue", "Priority", "Recommendation"],
    recs_content,
    col_widths=[2.0, 0.8, 3.9]
)

add_heading(doc, "Technical Gaps", 2)
recs_tech = [
    ("og:image is SVG", "Medium", "Most social platforms prefer PNG/JPEG for og:image. Convert social-preview.svg to a PNG for maximum compatibility."),
    ("Single HTML file", "Low", "As the site grows, maintaining all CSS/JS/HTML in one file becomes unwieldy. Consider extracting CSS and JS to separate files."),
    ("No form spam protection", "Medium", "Formspree provides basic spam filtering, but no honeypot or reCAPTCHA is implemented."),
    ("Login modal is a placeholder", "Low", "The Login button opens a \"Coming Soon\" modal. Implement actual authentication or remove the button."),
    ("No analytics", "Low", "No page view tracking. Consider privacy-respecting analytics (Plausible, Fathom) to understand visitor behavior."),
    ("No sitemap.xml or robots.txt", "Low", "Both would improve crawlability and indexation by search engines."),
    ("Form validation is client-side only", "Low", "Formspree provides server-side validation, but client errors should also display inline field messages rather than generic alerts."),
]
add_table(doc,
    ["Issue", "Priority", "Recommendation"],
    recs_tech,
    col_widths=[2.0, 0.8, 3.9]
)

# ── FOOTER ────────────────────────────────────────────────────────────────────

doc.add_page_break()
add_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Justice Innovations · Website Technical & Functional Report · "
    f"{datetime.date.today().strftime('%B %Y')} · Confidential"
)
set_font(run, size=9, color=(150, 150, 150))

# ── SAVE ──────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Saved: {OUT}")
